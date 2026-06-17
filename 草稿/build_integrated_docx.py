from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BODY_DIR = ROOT / "草稿" / "正文"
OUT_DIR = ROOT / "草稿" / "整合文档"
TITLE = "基于多模态深度学习的阻尼钢框架最大层间位移角预测研究"
OUTPUT = OUT_DIR / f"{TITLE}.docx"
FIGURE_DIR = ROOT / "草稿" / "论文插图" / "png"
FIGURES_BY_HEADING = {
    "1.3 本文的研究目标与主要贡献": [
        ("fig01_research_workflow.png", "图 1-1 研究技术路线"),
    ],
    "3.1 数据集来源与结构样本": [
        ("fig04_damped_frame_encoding.png", "图 3-1 阻尼钢框架参数化与布置编码"),
    ],
    "3.2 地震动输入与样本划分": [
        ("fig02_ground_motion_screening.png", "图 3-2 地震动筛选与调幅流程"),
    ],
    "3.3.1 地震动时频图": [
        ("fig05_input_representations.png", "图 3-4 输入特征表达方式"),
    ],
    "3.5 数据分布与研究边界": [
        ("fig08_dataset_distribution.png", "图 3-5 数据集划分与漂移区间分布"),
    ],
    "4.2 多模态 2D-CNN 主模型": [
        ("fig06_multimodal_2dcnn_architecture.png", "图 4-1 多模态 2D-CNN 主模型结构"),
    ],
    "4.3 时序深度学习模型": [
        ("fig07_sequence_baselines.png", "图 4-2 LSTM 与 WaveNet 时序基线模型"),
    ],
    "6.1 整体测试结果": [
        ("fig09_overall_model_performance.png", "图 6-1 各模型整体测试性能"),
    ],
    "6.4 分漂移区间误差分析": [
        ("fig10_tail_reliability.png", "图 6-2 中高漂移区间低估风险"),
    ],
}
FIGURES_AFTER_PARAGRAPH = [
    (
        "对每条地震动时程计算 Arias 强度累积曲线",
        "fig03_arias_window.png",
        "图 3-3 Arias 强度滑动窗口截取示意",
    ),
]
ENGLISH_ABSTRACT = (
    "Nonlinear time-history analysis can provide accurate evaluations of structural responses under earthquake "
    "excitation, yet its high computational cost limits its use in large-scale parametric analysis, rapid risk "
    "assessment, and preliminary scheme screening. The maximum interstory drift ratio is a key engineering demand "
    "parameter for characterizing global deformation and damage severity of structures during earthquakes, and it "
    "plays an essential role in performance-based seismic assessment and design. To enable rapid prediction of the "
    "maximum interstory drift ratio of damped steel frames, this study proposes a surrogate modeling approach that "
    "integrates ground-motion time-frequency images, structural parameters, ground-motion statistical features, and "
    "damper layout descriptors. A two-dimensional convolutional neural network (2D-CNN) incorporating building "
    "feature vectors is adopted as the primary model. To systematically evaluate the effectiveness of the proposed "
    "method, the primary model is compared with LSTM, WaveNet, MLP, Random Forest, XGBoost, LightGBM, and CatBoost "
    "under a unified data partitioning scheme and a consistent evaluation framework. The results show that the "
    "proposed model achieves the best overall prediction accuracy on the current test set across multiple metrics, "
    "indicating that the deep feature fusion mechanism can effectively improve response prediction in the low- to "
    "moderate-drift range. Further analysis reveals that, although the model performs well in terms of global "
    "metrics, systematic underestimation remains in the medium-to-high drift range, and its ability to identify "
    "high-drift thresholds is still limited. Overall, the proposed method can support rapid surrogate prediction of "
    "low- to moderate-drift responses of damped steel frames, but additional tail-region samples and improved "
    "training strategies are still required for high-drift risk identification and post-yield response prediction."
)
ENGLISH_KEYWORDS = (
    "Keywords: seismic response prediction; multimodal deep learning; damped steel frame; maximum interstory "
    "drift ratio; surrogate model; 2D-CNN"
)

CHAPTERS = [
    BODY_DIR / "1摘要正文.md",
    BODY_DIR / "0引言正文.md",
    BODY_DIR / "2_本文使用的深度学习模型_正文稿.md",
    BODY_DIR / "3_数据集与问题定义_正文稿.md",
    BODY_DIR / "4_方法_正文稿.md",
    BODY_DIR / "5_实验设计与评价指标_正文稿.md",
    BODY_DIR / "6_结果与分析_正文稿.md",
    BODY_DIR / "7_讨论_正文稿.md",
    BODY_DIR / "8_结论_正文稿.md",
]


def clean_title(text: str) -> str:
    text = text.strip()
    text = re.sub(r"[（(]正文稿[）)]", "", text)
    return text.strip()


def normalize_ref_key(text: str) -> str:
    text = re.sub(r"https?://\S+", "", text)
    text = re.sub(r"\bdoi\s*[:：]\s*\S+", "", text, flags=re.I)
    text = text.lower()
    text = text.replace("&", "and")
    text = re.sub(r"[*_`。．，,.;；:：()（）\[\]{}<>《》\"'“”‘’\\/\-–—]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:220]


def split_body_and_refs(text: str) -> tuple[str, dict[int, str]]:
    pattern = re.compile(r"^##\s*(?:本章)?参考文献.*$", re.M)
    match = pattern.search(text)
    if not match:
        return text.strip(), {}

    body = text[: match.start()].strip()
    refs_text = text[match.end() :].strip()
    refs: dict[int, str] = {}
    current_num: int | None = None
    current_parts: list[str] = []

    for raw_line in refs_text.splitlines():
        line = raw_line.strip()
        if not line or line.startswith(">"):
            continue
        ref_match = re.match(r"^\[(\d+)\]\s*(.+)$", line)
        if ref_match:
            if current_num is not None and current_parts:
                refs[current_num] = " ".join(current_parts).strip()
            current_num = int(ref_match.group(1))
            current_parts = [ref_match.group(2).strip()]
        elif current_num is not None:
            current_parts.append(line)
    if current_num is not None and current_parts:
        refs[current_num] = " ".join(current_parts).strip()

    return body, refs


def parse_citation_numbers(citation: str) -> list[int] | None:
    nums: list[int] = []
    for part in citation.split(","):
        part = part.strip()
        if not part:
            return None
        if "-" in part:
            bits = [b.strip() for b in part.split("-", 1)]
            if len(bits) != 2 or not all(b.isdigit() for b in bits):
                return None
            start, end = int(bits[0]), int(bits[1])
            if start > end:
                return None
            nums.extend(range(start, end + 1))
        elif part.isdigit():
            nums.append(int(part))
        else:
            return None
    return nums


def format_citation(nums: list[int]) -> str:
    unique: list[int] = []
    for n in nums:
        if n not in unique:
            unique.append(n)

    runs: list[str] = []
    i = 0
    while i < len(unique):
        start = unique[i]
        end = start
        j = i + 1
        while j < len(unique) and unique[j] == end + 1:
            end = unique[j]
            j += 1
        if end > start:
            runs.append(f"{start}-{end}")
        else:
            runs.append(str(start))
        i = j
    return "[" + ",".join(runs) + "]"


def replace_citations(body: str, mapping: dict[int, int]) -> str:
    citation_re = re.compile(r"\[((?:\d+\s*(?:-\s*\d+)?)(?:\s*,\s*\d+\s*(?:-\s*\d+)?)*)\]")

    def repl(match: re.Match[str]) -> str:
        nums = parse_citation_numbers(match.group(1))
        if not nums or any(n not in mapping for n in nums):
            return match.group(0)
        return format_citation([mapping[n] for n in nums])

    return citation_re.sub(repl, body)


GREEK = {
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "delta": "δ",
    "epsilon": "ε",
    "theta": "θ",
    "lambda": "λ",
    "mu": "μ",
    "phi": "φ",
    "sigma": "σ",
    "tau": "τ",
}


def latex_to_text(text: str) -> str:
    text = text.replace("\\(", "").replace("\\)", "")
    text = text.replace("\\[", "").replace("\\]", "")
    text = text.replace("\\left", "").replace("\\right", "")
    text = text.replace("\\quad", " ")
    text = text.replace("\\,", " ")
    text = text.replace("\\ ", " ")
    text = text.replace("*_{", "∗_{")

    # Convert simple LaTeX wrappers while preserving the mathematical symbol.
    wrapper = re.compile(r"\\(?:mathbf|mathcal|boldsymbol|mathrm|mathbb|operatorname)\{([^{}]+)\}")
    while wrapper.search(text):
        text = wrapper.sub(lambda m: latex_to_text(m.group(1)), text)

    accent_map = {"hat": "\u0302", "bar": "\u0304", "tilde": "\u0303"}
    for command, accent in accent_map.items():
        pattern = re.compile(rf"\\{command}\{{([^{{}}]+)\}}")
        while pattern.search(text):
            text = pattern.sub(lambda m, a=accent: latex_to_text(m.group(1)) + a, text)

    frac = re.compile(r"\\frac\{([^{}]+)\}\{([^{}]+)\}")
    while frac.search(text):
        text = frac.sub(lambda m: f"({latex_to_text(m.group(1))})/({latex_to_text(m.group(2))})", text)

    sqrt = re.compile(r"\\sqrt\{([^{}]+)\}")
    text = sqrt.sub(lambda m: f"√({latex_to_text(m.group(1))})", text)

    replacements = {
        "\\sum": "Σ",
        "\\max": "max",
        "\\min": "min",
        "\\tanh": "tanh",
        "\\ln": "ln",
        "\\exp": "exp",
        "\\odot": "⊙",
        "\\cdot": "·",
        "\\times": "×",
        "\\leq": "≤",
        "\\le": "≤",
        "\\geq": "≥",
        "\\ge": "≥",
        "\\%": "%",
        "\\_": "_",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    for name, symbol in GREEK.items():
        text = re.sub(rf"\\{name}\b", symbol, text)

    text = re.sub(r"\\([A-Za-z]+)", r"\1", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def convert_inline_math(text: str) -> str:
    return re.sub(r"\\\((.+?)\\\)", lambda m: f"${m.group(1).strip()}$", text)


def strip_markdown_outside_math(text: str) -> str:
    parts = re.split(r"(\$[^$]*\$)", text)
    cleaned: list[str] = []
    for idx, part in enumerate(parts):
        if idx % 2 == 1:
            cleaned.append(part)
        else:
            part = part.replace("**", "").replace("__", "")
            part = part.replace("*", "")
            part = part.replace("`", "")
            cleaned.append(part)
    return "".join(cleaned)


def clean_inline(text: str) -> str:
    text = convert_inline_math(text)
    text = strip_markdown_outside_math(text)
    text = text.replace("&amp;", "&")
    return text.strip()


def set_run_font(run, size: float | None = None, bold: bool | None = None, east_asia: str = "宋体"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_runs(paragraph, text: str, size: float = 11, bold_default: bool = False):
    text = clean_inline(text)
    if not text:
        return
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        if is_bold:
            part = part[2:-2]
        run = paragraph.add_run(part)
        set_run_font(run, size=size, bold=bold_default or is_bold)


def set_paragraph_format(paragraph, before=0, after=6, line=1.1, justify=True):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int, col_widths: list[int]):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in col_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, col_widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [clean_inline(c.strip()) for c in raw.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    for row in rows:
        row.extend([""] * (col_count - len(row)))

    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    total = 9360
    # Give the first descriptive column a little more room for engineering tables.
    if col_count == 2:
        widths = [2600, total - 2600]
    elif col_count == 3:
        widths = [2200, 3580, 3580]
    elif col_count == 4:
        widths = [2400, 2320, 2320, 2320]
    elif col_count == 5:
        widths = [2200, 1790, 1790, 1790, 1790]
    elif col_count == 6:
        widths = [1900, 1492, 1492, 1492, 1492, 1492]
    elif col_count == 7:
        widths = [1700, 1276, 1276, 1276, 1276, 1276, 1280]
    elif col_count == 8:
        widths = [1600] + [1109] * 7
        widths[-1] += total - sum(widths)
    else:
        widths = [total // col_count] * col_count
        widths[-1] += total - sum(widths)
    set_table_width(table, total, widths)

    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                shade_cell(cell, "F2F4F7")
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 or c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_runs(p, text, size=8.0, bold_default=(r_idx == 0))

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(6)


def add_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    add_runs(paragraph, clean_title(text), size={1: 16, 2: 13, 3: 12}.get(level, 12), bold_default=True)
    if level == 1:
        paragraph.paragraph_format.page_break_before = text.strip().startswith(tuple(str(i) for i in range(1, 9)))


def add_paragraph(doc: Document, text: str, style: str | None = None, indent: bool = True):
    paragraph = doc.add_paragraph(style=style)
    set_paragraph_format(paragraph)
    if indent:
        paragraph.paragraph_format.first_line_indent = Inches(0.28)
    add_runs(paragraph, text)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_runs(p, text, size=10, bold_default=True)


def add_figure(doc: Document, image_path: Path, caption: str, width_in: float = 6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    add_caption(doc, caption)


def add_paper_figure(doc: Document, filename: str, caption: str, width_in: float = 6.5):
    image_path = FIGURE_DIR / filename
    if image_path.exists():
        add_figure(doc, image_path, caption, width_in=width_in)


def add_display_formula(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    formula = " ".join(line.strip() for line in lines if line.strip())
    run = p.add_run(f"$$ {formula} $$")
    set_run_font(run, size=10.5, bold=False)


def add_markdown_body(doc: Document, body: str, skip_first_heading: bool = False):
    lines = body.splitlines()
    i = 0
    in_equation = False
    equation_lines: list[str] = []
    first_heading_skipped = False
    inserted_after_paragraph: set[str] = set()

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if not line:
            i += 1
            continue

        if line in {"\\[", "$$"}:
            in_equation = True
            equation_lines = []
            i += 1
            continue
        if in_equation:
            if line in {"\\]", "$$"}:
                add_display_formula(doc, equation_lines)
                in_equation = False
            else:
                equation_lines.append(line)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_markdown_table(lines, i)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            if skip_first_heading and not first_heading_skipped:
                first_heading_skipped = True
            else:
                heading_text = clean_title(heading.group(2))
                add_heading(doc, heading_text, len(heading.group(1)))
                for filename, caption in FIGURES_BY_HEADING.get(heading_text, []):
                    add_paper_figure(doc, filename, caption)
            i += 1
            continue

        if re.match(r"^(表|图)\s*\d", line):
            add_caption(doc, line)
            i += 1
            continue

        bullet = re.match(r"^[-*+]\s+(.+)$", line)
        if bullet:
            add_paragraph(doc, bullet.group(1), style="List Bullet", indent=False)
            i += 1
            continue

        add_paragraph(doc, line)
        for trigger, filename, caption in FIGURES_AFTER_PARAGRAPH:
            if trigger in line and filename not in inserted_after_paragraph:
                add_paper_figure(doc, filename, caption)
                inserted_after_paragraph.add(filename)
        i += 1


def apply_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def build_reference_maps():
    global_refs: list[str] = []
    keys: dict[str, int] = {}
    processed: list[tuple[Path, str, dict[int, int]]] = []

    for path in CHAPTERS:
        body, refs = split_body_and_refs(path.read_text(encoding="utf-8-sig"))
        local_to_global: dict[int, int] = {}
        for local_num in sorted(refs):
            ref_text = clean_inline(refs[local_num])
            key = normalize_ref_key(ref_text)
            if key in keys:
                global_num = keys[key]
            else:
                global_refs.append(ref_text)
                global_num = len(global_refs)
                keys[key] = global_num
            local_to_global[local_num] = global_num
        body = replace_citations(body, local_to_global)
        processed.append((path, body, local_to_global))
    return processed, global_refs


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(170)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(TITLE)
    set_run_font(run, size=22, bold=True, east_asia="黑体")
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(120)
    run = p2.add_run("正文整合稿")
    set_run_font(run, size=14, bold=False, east_asia="宋体")

    doc.add_page_break()


def add_references(doc: Document, refs: list[str]):
    add_heading(doc, "参考文献", 1)
    for idx, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.24)
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.05
        add_runs(p, f"[{idx}] {ref}", size=10)


def add_english_abstract(doc: Document):
    add_heading(doc, "Abstract", 1)
    add_paragraph(doc, ENGLISH_ABSTRACT, indent=False)
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=4, after=6, line=1.1, justify=True)
    add_runs(paragraph, ENGLISH_KEYWORDS, size=11)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    processed, refs = build_reference_maps()

    doc = Document()
    apply_styles(doc)
    add_title_page(doc)

    for path, body, _ in processed:
        if path.name == "1摘要正文.md":
            add_heading(doc, "摘要", 1)
            add_markdown_body(doc, body, skip_first_heading=True)
            add_english_abstract(doc)
            doc.add_page_break()
        else:
            add_markdown_body(doc, body)

    add_references(doc, refs)
    doc.save(OUTPUT)
    print(OUTPUT)
    print(f"references={len(refs)}")


if __name__ == "__main__":
    main()
